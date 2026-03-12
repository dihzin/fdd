from __future__ import annotations

import uuid
from datetime import UTC, datetime
from pathlib import Path

from docx import Document as WordDocument

from app.common.enums import DocumentStatus, DocumentType, SectionStatus
from app.domains.documents.docx_export_service import FddDocxExportService
from app.domains.documents.models import DocumentSection
from app.domains.documents.schemas import DocumentUpdate


class FakeTemplate:
    def __init__(self, storage_path: str) -> None:
        self.storage_path = storage_path


class FakeProject:
    def __init__(self) -> None:
        self.id = uuid.uuid4()
        self.code = "PRJ-001"
        self.name = "SAP OTC Rollout"
        self.client_name = "Contoso"
        self.description = "Project description"


class FakeSolution:
    def __init__(self) -> None:
        self.id = uuid.uuid4()
        self.name = "Billing Localization"
        self.description = "Solution desc"
        self.module = "SD"
        self.phase = "Design"
        self.status = type("Status", (), {"value": "draft"})()
        self.wizard_context = {}
        self.solution_requirements = []


class FakeSection:
    def __init__(self, section_key: str, content_text: str) -> None:
        self.id = uuid.uuid4()
        self.section_key = section_key
        self.section_title = section_key
        self.section_order = 1
        self.content_text = content_text
        self.content_json = {}
        self.status = SectionStatus.DRAFT
        self.updated_by_id = None
        self.updated_at = datetime.now(UTC)


class FakeDocument:
    def __init__(self, template_path: str) -> None:
        self.id = uuid.uuid4()
        self.project_id = uuid.uuid4()
        self.solution_id = uuid.uuid4()
        self.template_id = uuid.uuid4()
        self.title = "FDD Export"
        self.document_type = DocumentType.FDD
        self.status = DocumentStatus.DRAFT
        self.version = 1
        self.generated_file_path = None
        self.snapshot_payload = None
        self.created_by_id = None
        self.approved_by_id = None
        self.approved_at = None
        self.created_at = datetime.now(UTC)
        self.updated_at = datetime.now(UTC)
        self.project = FakeProject()
        self.solution = FakeSolution()
        self.template = FakeTemplate(template_path)
        self.sections = [
            FakeSection("objective_scope", "Objective section content"),
            FakeSection("business_process", "Business process content"),
            FakeSection("requirements_coverage", "Requirements coverage content"),
            FakeSection("functional_design", "Functional design content"),
            FakeSection("interfaces_data", "Interfaces data content"),
            FakeSection("security_controls", "Security controls content"),
            FakeSection("assumptions_dependencies", "Assumptions content"),
            FakeSection("test_considerations", "Test considerations content"),
        ]


class FakeDocumentRepository:
    def __init__(self, document: FakeDocument) -> None:
        self.document = document

    def get(self, document_id: uuid.UUID):
        return self.document if self.document.id == document_id else None

    def update(self, document: FakeDocument, payload: DocumentUpdate):
        data = payload.model_dump(exclude_unset=True)
        for field, value in data.items():
            setattr(document, field, value)
        return document


def test_docx_export_renders_template_and_updates_document_path(tmp_path: Path, monkeypatch):
    template_path = tmp_path / "template.docx"
    output_dir = tmp_path / "exports"
    output_dir.mkdir()

    template = WordDocument()
    template.add_heading("{{ document_title }}", 0)
    template.add_paragraph("{{ fdd_objective_scope }}")
    template.add_paragraph("{{ fdd_business_process }}")
    template.save(template_path)

    document = FakeDocument(str(template_path))
    repository = FakeDocumentRepository(document)

    monkeypatch.setattr("app.domains.documents.docx_export_service.settings.document_export_dir", output_dir)

    service = FddDocxExportService(repository)
    exported = service.export(document.id)

    assert exported.output_path.exists()
    assert document.generated_file_path == str(exported.output_path)

    rendered = WordDocument(exported.output_path)
    full_text = "\n".join(paragraph.text for paragraph in rendered.paragraphs)
    assert "FDD Export" in full_text
    assert "Objective section content" in full_text
